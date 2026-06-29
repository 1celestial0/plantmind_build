#!/usr/bin/env python3
"""Generate ultra implementation + win strategy Word doc for technical team handover."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "deliverables"
NAVY = RGBColor(0x06, 0x5A, 0x82)
TEAL = RGBColor(0x1C, 0x72, 0x93)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY if level == 1 else TEAL
    return p


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].text = hd
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(v)
    doc.add_paragraph()


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PlantMind × Götze Engine\n")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    r2 = p.add_run(
        "\nUltra Implementation & Integration Guide\n"
        "For Technical Team · TL · Delivery Manager · Reviewers\n"
        "Version 1.0 · 29 June 2026"
    )
    r2.font.size = Pt(14)
    doc.add_page_break()

    h(doc, "PART 0 — What Is This Project?", 1)
    doc.add_paragraph(
        "PlantMind converts industrial sensor data into ONE ranked, explainable, "
        "human-approved maintenance action — and proves that action would rescue the asset."
    )
    tbl(doc, ["Question", "Answer"], [
        ("What is it?", "Physics-Informed Engineering Intelligence + Götze Decision Engine"),
        ("What is it NOT?", "Not a chatbot, not a dashboard-only tool, not three projects"),
        ("Workspace", r"C:\Users\hp\Claude\Projects\PlantMind-Live"),
        ("Event", "LTTS Global EI Hackathon · 9 July 2026 · 24h · 4 members"),
        ("Tagline", "Predict the failure. Decide the fix. Prove it."),
    ])

    h(doc, "Two Implementations — One Product", 2)
    tbl(doc, ["", "v1 forge-v1 (BUILT)", "v2 src/ (TARGET)"], [
        ("Location", "src/legacy/forge-v1/", "src/"),
        ("Architecture", "5-layer MetaGPT", "5-agent LangGraph"),
        ("Scoring", "G-score", "IIS (canonical)"),
        ("Health", "RandomForest C-MAPSS", "Weibull multi-asset"),
        ("Hackathon", "SHIP THIS if time short", "Build toward post-demo"),
    ])
    doc.add_paragraph("Rule: LOCKED_STATE.md is product truth. forge-v1 is insurance demo.")

    h(doc, "PART 1 — Codebase Lineage", 1)
    h(doc, "v1 Data Lineage (Built)", 2)
    flow1 = (
        "train_FD001.txt → ingestion.py → features.py → model.py (RF RUL) → "
        "EngineHealth → gotze_engine.py (G-score) → proof_engineer.py → app.py (4 tabs)"
    )
    doc.add_paragraph(flow1)
    tbl(doc, ["File", "Input", "Output"], [
        ("ingestion.py", "CSV path", "DataFrame + RUL"),
        ("features.py", "DataFrame", "Feature matrix"),
        ("model.py", "Features", "RUL float"),
        ("gotze_engine.py", "EngineHealth", "EngineDecision + proof"),
        ("pipeline.py", "engine_id", "decision + chart + trace"),
        ("app.py", "pipeline", "Streamlit UI"),
    ])

    h(doc, "v2 Data Lineage (Target)", 2)
    flow2 = (
        "synthetic/Kaggle → ingest → features → weibull.py → DataSentinel → "
        "AssetHealthOracle → [TRIGGER] → GötzeEngine (IIS) → RootCauseAnalyst → "
        "ExecutiveSummarizer → FastAPI JSON → dashboard → Human Approve → audit"
    )
    doc.add_paragraph(flow2)

    h(doc, "v1 → v2 Migration Map", 2)
    tbl(doc, ["v1 Module", "v2 Destination", "Action"], [
        ("ingestion.py", "ml/data + src/physics", "Port + extend"),
        ("gotze_engine.py", "src/agents/gotze_engine.py", "REWRITE G→IIS"),
        ("pipeline.py", "src/pipeline/orchestrator.py", "REWRITE LangGraph"),
        ("app.py", "src/dashboard/app.py", "Add approve + IIS panel"),
        ("messages.py", "src/contracts/", "Extend schemas"),
    ])

    h(doc, "PART 2 — Phase-by-Phase Build (P0–P6)", 1)
    tbl(doc, ["Phase", "Name", "Owner", "Deliverable", "Test Gate"], [
        ("P0", "Contracts + env", "Lane 1", "src/contracts/*.py", "pytest contracts"),
        ("P1", "Physics + synthetic", "Lane 2", "weibull.py + generate_data.py", "health<40 pump_07"),
        ("P2", "Agents 1-2", "Lane 1", "Sentinel + Oracle", "anomaly + health tests"),
        ("P3", "Götze + IIS", "Lane 1", "gotze_engine.py IIS", "correct winner Scenario A"),
        ("P4", "Agents 4-5 + API", "Lane 1", "FastAPI + RAG + audit", "curl /evaluate 200"),
        ("P5", "Dashboard + approve", "Lane 3", "Streamlit v2", "approve writes audit"),
        ("P6", "E2E + freeze", "Lane 5", "tag + backup video", "12-test matrix pass"),
    ])

    for phase, title, why, how, who, files, test in [
        ("P0", "Environment & Contracts",
         "Without frozen contracts, lanes build incompatible modules.",
         "Create Pydantic models matching LOCKED_STATE §4. pip install requirements.",
         "Lane 1 — Sourav",
         "src/contracts/physics.py, agents.py, governance.py, state.py",
         "pytest tests/test_contracts.py -v"),
        ("P1", "Physics & Synthetic Data",
         "Oracle needs Weibull health. Synthetic data enables live failure injection.",
         "ml/synthesis/generate_data.py; src/physics/weibull.py; calibrate from CMAPSS.",
         "Lane 2 — Sourav",
         "H(t)=100·exp(−λ·S·t^β); 30 assets; plant_config.yaml",
         "pump_07 cycle 400 → health<40, rul_days<14"),
        ("P2", "DataSentinel + AssetHealthOracle",
         "Feed Götze trigger. Sentinel flags only; Oracle always returns ci_95.",
         "Z-score + Mahalanobis; call PhysicsModelInterface.",
         "Lane 1 — Sourav",
         "src/agents/data_sentinel.py, asset_health_oracle.py",
         "Inject z>3 → critical; health report valid JSON"),
        ("P3", "GötzeEngine + IIS",
         "THE PRODUCT. Scores all interventions → ONE winner.",
         "IIS = 0.35·ΔP + 0.25·ΔCost + 0.20·Feas + 0.15·Hist − 0.05·Safety",
         "Lane 1 — Sourav",
         "src/agents/gotze_engine.py; interventions.yaml",
         "Scenario A → reduce_load wins; safety veto works"),
        ("P4", "RootCause + Summarizer + FastAPI",
         "Agentic story for judges: citations + brief + API for UI.",
         "ChromaDB RAG; LangGraph orchestrator; audit hash chain.",
         "Lane 1 — Sourav",
         "src/api/main.py; src/pipeline/orchestrator.py",
         "POST /evaluate returns full agent JSON"),
        ("P5", "Dashboard + Human Approve",
         "Judges touch the product. Governance proof.",
         "Streamlit reads API only; approve → POST /approve.",
         "Lane 3 — Members 2/3",
         "src/dashboard/app.py; iis_panel; proof_chart from forge-v1",
         "Manual Scenario A + approve/reject"),
        ("P6", "Integration & Demo Freeze",
         "Won on stage not in docs.",
         "E2E tests; scenario injector; git tag v1.0-hackathon-submission.",
         "Lane 5 — Member 4",
         "tests/test_e2e_scenario_a.py; backup video",
         "12-test matrix; 5-min script rehearsed 5+ times"),
    ]:
        h(doc, f"{phase} — {title}", 2)
        doc.add_paragraph(f"WHY: {why}")
        doc.add_paragraph(f"HOW: {how}")
        doc.add_paragraph(f"WHO: {who}")
        doc.add_paragraph(f"FILES: {files}")
        doc.add_paragraph(f"TEST: {test}")

    h(doc, "PART 3 — Integration & Test Matrix", 1)
    tbl(doc, ["Test ID", "What", "Pass"], [
        ("T01", "Contracts JSON", "pytest test_contracts"),
        ("T02", "Weibull health", "pump_07 health<40"),
        ("T03", "Sentinel", "z>3 → critical"),
        ("T04", "IIS winner", "Scenario A correct action"),
        ("T05", "Safety veto", "unsafe blocked"),
        ("T06", "API evaluate", "HTTP 200 + full JSON"),
        ("T07", "Approve flow", "audit written"),
        ("T08", "RAG citation", "citation or uncertain"),
        ("T09", "LLM fallback", "templates when Groq down"),
        ("T10", "v1 fallback", "forge-v1 app starts"),
        ("T11", "Scenario B", "emergency_stop wins"),
        ("T12", "Scenario D", "sensor dropout = data issue"),
    ])

    h(doc, "PART 4 — Team Lane Rules", 1)
    tbl(doc, ["Lane", "Member", "WRITE folders", "READ only"], [
        ("1", "Sourav", "src/agents, api, pipeline, governance", "contracts, physics"),
        ("2", "Sourav", "src/physics, ml/", "contracts"),
        ("3", "M2/M3", "src/dashboard/", "contracts JSON"),
        ("4", "Team", "deploy/databricks/", "architecture"),
        ("5", "M4", "ops/runbooks, tests/", "docs only"),
    ])
    doc.add_paragraph("PR rule: never merge contract changes without LOCKED_STATE vault update.")

    h(doc, "PART 5 — Hackathon Decision Tree", 1)
    doc.add_paragraph("July 8+ → demo forge-v1 ONLY, freeze, rehearse.")
    doc.add_paragraph("P3 not done → build v2 OR prep forge-v1 demo in parallel.")
    doc.add_paragraph("P5 not done → demo forge-v1 + narrate v2 governance as production roadmap.")
    doc.add_paragraph("P5 done → demo v2 primary, forge-v1 as backup video.")

    h(doc, "PART 6 — Win Strategy (Honest)", 1)
    doc.add_paragraph("There is NO 200% win guarantee. Estimated composite score today: 7.8/10.")
    doc.add_paragraph("With polished forge-v1 demo + tight pitch: 8.5/10 — top-tier candidate.")
    tbl(doc, ["Dimension", "Score", "Action"], [
        ("Strategic alignment", "8.5/10", "Name LTTS client vertical"),
        ("Story", "9.0/10", "Rehearse Götze hook 20x"),
        ("Demo (built)", "7.5/10", "Confirm live run + backup video"),
        ("Governance", "6.0 v1 / 8.5 v2", "Approve button or honest prod slide"),
        ("Execution risk", "6.5/10", "ONE folder, ONE demo path"),
    ])
    doc.add_paragraph(
        "You chose the RIGHT project for LTTS EI hackathon. Risk is execution clarity, not the idea."
    )
    doc.add_paragraph("Maximum win path: PlantMind-Live only → demo forge-v1 → pitch v2 vision → rehearse.")

    h(doc, "PART 7 — Multi-AI Operating System", 1)
    tbl(doc, ["Tool", "Entry file", "Session start"], [
        ("Claude Code", "CLAUDE.md", "Read AI-OPERATING-SYSTEM.md"),
        ("Grok CLI", "AGENTS.md", "Same + declare lane"),
        ("Gemini CLI", "GEMINI.md", "Same + declare lane"),
        ("Human", "00-START-HERE.md", "scripts/start-session.ps1"),
    ])
    doc.add_paragraph("Archives (DO NOT WRITE): PlantMind/, PlantMind_hckthn/")
    doc.add_paragraph("Close every session: 'close session' → ROADMAP + Chat Context + git commit")

    h(doc, "Appendix — Commands", 1)
    cmds = [
        r'cd "C:\Users\hp\Claude\Projects\PlantMind-Live"',
        r".\scripts\start-session.ps1",
        r"streamlit run src\legacy\forge-v1\app.py",
        r"uvicorn src.api.main:app --reload",
        r"pytest tests/ -v",
    ]
    for c in cmds:
        doc.add_paragraph(c, style="List Bullet")

    out = OUT / "PlantMind_Ultra_Implementation_Team_Guide.docx"
    doc.save(str(out))
    print(f"Ultra team doc: {out}")
    return out


if __name__ == "__main__":
    build()