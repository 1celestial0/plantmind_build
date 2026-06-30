"""
gen_team_docs.py — regenerate the standard team deliverable set from the locked truth.
Outputs (to team-share/): Executive Brief (DOCX) + Feature/Rubric/Top-20 Tracker (XLSX).
Source of truth: PROJECT-DNA.md v1.0 + LOCKED_STATE.md. Re-run after any DNA amendment.
Usage:  python scripts/gen_team_docs.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "team-share"
OUT.mkdir(exist_ok=True)
NAVY = "003366"; ORANGE = "FF6B00"

# ---------- data (mirrors PROJECT-DNA §6/§7 + ROADMAP Top-20) ----------
FEATURES = [
    ("F-01","Götze Decision Panel","P1 · agent","BUILDING","L1+L3","Select degraded asset → top action + gap + Approve"),
    ("F-02","Plant Config Manifest + Viewer","P2 · L0","BUILDING (high risk)","L1+L3","Viewer shows manifest driving pipeline + profile swap"),
    ("F-03","MaintenanceScheduler","P1 · agent","BUILDING","L1","Approve → work order + audit entry"),
    ("F-04","DataSentinel","P1 · agent","BUILDING","L1","Inject anomaly → typed flag + severity"),
    ("F-05","AssetHealthOracle","P1 · physics","BUILDING","L2","Health curve degrades; RUL+CI render (§6a λ/β)"),
    ("F-06","RootCauseAnalyst","P1 · agent","BUILDING","L1","'Why?' expands to cited manual/fault excerpts"),
    ("F-07","ExecutiveSummarizer","P1 · agent","BUILDING","L1","3-bullet brief + $180k figure"),
    ("F-08","IIS Scoring Engine & Profiles","P1+P2 · logic","BUILDING","L1","Swap profile → recommendation reorders"),
    ("F-09","Governance: Audit/Hash-Chain/Lineage","P1 · governance","BUILDING (verify)","L1+L3","Audit explorer shows chain valid + full state"),
    ("F-10","Two-Layer Contracts (Layer-0 IP)","P2 · L0","BUILDING — TOP LOCK","L1","Manifest + WorkOrder models validate"),
    ("F-11","Synthetic Data + Injector","P1/P2 · data","BUILDING","L2/L5","Inject → PUMP-001 health drops → Götze fires"),
    ("F-12","Databricks Medallion (Layer 1)","P2 · L1","BUILDING (parallel)","L4","Unity Catalog lineage Bronze→Gold shown"),
    ("F-13","Plant Overview Dashboard","P1 · UI","BUILDING (reskin)","L3","Grid loads; hero red trigger; click → detail"),
    ("F-14","Asset Detail Hero Screen","P1 · UI","BUILDING (reskin)","L3","PUMP-001 → gauge + trends + Götze panel"),
    ("F-15","Orchestrator State Machine","P1+P2 · core","BUILDING","L1","Kill one agent → continues; panel still renders"),
]
RUBRIC = [
    (1,"Innovation / novelty",15,4.0,12.0,"How is this not just a digital twin / agent wrapper?"),
    (2,"Technical depth & feasibility",12,3.0,7.2,"Show the config composing the pipeline, live."),
    (3,"Business impact / ROI",18,4.0,14.4,"Defend the $180k line by line."),
    (4,"Fit to existing data / zero rip-replace",10,3.0,6.0,"Show it with MY historian's tag names."),
    (5,"Scalability / modularity",13,3.5,9.1,"Onboard a new asset class now, no code."),
    (6,"Governance / trust / explainability",10,3.5,7.0,"Prove this decision wasn't tampered with."),
    (7,"Demo quality / storytelling",10,2.5,5.0,"Run the whole thing end-to-end, no cuts."),
    (8,"Databricks / partnership alignment",12,4.0,9.6,"Why does this need Databricks specifically?"),
]
TOP20 = [
    (1,"DNA ratified → v1.0 LOCKED","DONE","P0"),(2,"Lock src/contracts: manifest.py + workorder.py","TODO","P0 🔴"),
    (3,"Doc consolidation + 2 vaults + archive","DONE","P0"),(4,"Wire session-start to PROJECT-DNA","DONE","P0"),
    (5,"Continuity commit / first GitHub push","DONE","P0"),(6,"config/plants/hero.yaml (§6a λ/β)","TODO","P1"),
    (7,"Fix λ/β + RUL traps in src/physics","TODO","P1 ⚠️"),(8,"Start Databricks trial + Bronze","TODO","P1 🔴"),
    (9,"Config loader/composer (manifest→pipeline)","TODO","P2 ⚠️"),(10,"Hero pump end-to-end via manifest","TODO","P2"),
    (11,"IIS profile swap reliability→energy","TODO","P2"),(12,"Keep static fallback path runnable","TODO","P2"),
    (13,"Build MaintenanceScheduler agent","TODO","P3"),(14,"Full 6-agent orchestrator e2e","TODO","P3"),
    (15,"RAG corpus 10-20 manuals","TODO","P3"),(16,"UI dark→light reskin + Plant Overview","TODO","P4"),
    (17,"Config Viewer + Asset Detail + Audit view","TODO","P4"),(18,"Generate Spec/Scorecard/PPT from DNA","TODO","P5"),
    (19,"Demo flow + $180k + Lane 5 script","TODO","P5"),(20,"Rehearsal + freeze (H-14/16/20)","TODO","P6"),
]
CLOSED = [
    ("C1","Headline innovation","Two co-equal pillars (P1+P2)"),("C2","6th agent","MaintenanceScheduler (SafetyGuardian rejected)"),
    ("C3","UI theme","Light palette #FAFAFA/#003366/#FF6B00"),("C4","Manifest scope","Full config pipeline + Hour-16 fallback"),
    ("C5","Physics","Weibull always; PINN stretch (freeze H-14)"),("C6","IIS weights","Fixed in demo"),
    ("C7","Data","Synthetic, physics-seeded — by design"),("C8","Demo shape","3 assets · PUMP-001 · ~$180k · 1 swap"),
    ("C9","Demo platform","Local spine; Databricks live only if stable H-20"),("C10","LLM fallback","Pre-cached Götze narrative"),
]

# ---------- Excel tracker ----------
def excel():
    wb = Workbook()
    hf = Font(bold=True, color="FFFFFF"); hfill = PatternFill("solid", fgColor=NAVY)
    thin = Border(*[Side(style="thin", color="DDDDDD")]*4)
    def style_header(ws, ncol):
        for c in range(1, ncol+1):
            cell = ws.cell(1, c); cell.font = hf; cell.fill = hfill
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    def body(ws):
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True); cell.border = thin

    ws = wb.active; ws.title = "Features"
    ws.append(["ID","Feature","Category","Status","Lane","Demo proof"])
    for r in FEATURES: ws.append(list(r))
    for col,w in zip("ABCDEF",[7,32,16,20,10,46]): ws.column_dimensions[col].width = w
    style_header(ws,6); body(ws)

    ws2 = wb.create_sheet("Rubric")
    ws2.append(["#","Dimension","Weight","Maturity (1-5)","Weighted","Judge's killer question"])
    for r in RUBRIC: ws2.append(list(r))
    ws2.append(["","TOTAL READINESS",100,"", round(sum(r[4] for r in RUBRIC),1),"Target ≥90% by Jul 8"])
    tot = ws2.max_row
    for c in range(1,7): ws2.cell(tot,c).font = Font(bold=True, color=ORANGE)
    for col,w in zip("ABCDEF",[5,34,9,14,11,46]): ws2.column_dimensions[col].width = w
    style_header(ws2,6); body(ws2)

    ws3 = wb.create_sheet("Top-20")
    ws3.append(["#","Goal","Status","Phase"])
    for r in TOP20: ws3.append(list(r))
    for col,w in zip("ABCD",[5,46,8,8]): ws3.column_dimensions[col].width = w
    style_header(ws3,4); body(ws3)

    ws4 = wb.create_sheet("Closed-Decisions")
    ws4.append(["#","Question","Locked answer"])
    for r in CLOSED: ws4.append(list(r))
    for col,w in zip("ABC",[6,22,52]): ws4.column_dimensions[col].width = w
    style_header(ws4,3); body(ws4)

    p = OUT / "PlantMind_Tracker_v1.xlsx"; wb.save(p); return p

# ---------- DOCX executive brief ----------
def docx():
    d = Document()
    def H(text, size=14, color=NAVY, after=6):
        p = d.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color); p.paragraph_format.space_after = Pt(after); return p
    H("PlantMind × Götze Engine — Executive Brief (v1)", 18)
    sub = d.add_paragraph(); sr = sub.add_run("Derived from PROJECT-DNA v1.0 (LOCKED) · 2026-07-01 · for leadership"); sr.italic=True; sr.font.size=Pt(9)
    H("What it is", 13)
    d.add_paragraph("A config-driven, physics-informed, agentic decision fabric for asset-intensive industries. It turns existing "
                    "plant data into trusted, ranked, human-approved, audited engineering actions — without rip-and-replace. "
                    "It is NOT a digital twin, an alerting tool, or a generic agent builder; it is the governed decision-and-action layer on top of them.")
    H("The problem", 13)
    d.add_paragraph("Plants drown in alerts and tribal knowledge but lack a governed, auditable way to turn data into ranked, feasible "
                    "actions — and every new plant needs custom integration. Pilots don't scale.")
    H("The solution — two pillars", 13)
    d.add_paragraph("1) Closed decision loop: physics health → Intervention Impact Score → one approved, audited action.\n"
                    "2) Config-driven modularity: a Plant Config Manifest composes the whole stack; a new plant/asset is a config change, not a code deploy.")
    H("Proof & impact", 13)
    d.add_paragraph("Demo: a 30-asset fleet; inject degradation on hero PUMP-001 → health drops → the Götze panel surfaces the one best "
                    "action → human approves → work order + immutable audit. Headline: ~$180k downtime saved per prevented pump failure.")
    H("Why LTTS wins", 13)
    d.add_paragraph("A licensable Layer-0 Center-of-Excellence framework; swappable scoring profiles map 1:1 to the five LTTS–Databricks "
                    "joint areas; accelerates the Databricks go-to-market as the decision-fabric layer.")
    H("Status & ask", 13)
    d.add_paragraph("Product idea LOCKED (PROJECT-DNA v1.0). Readiness 70.3% against an 8-dimension rubric; gaps in demo rehearsal, "
                    "config-pipeline feasibility, and data-fit. Next blocker: lock the shared contracts. Hackathon: 9 July 2026.")
    foot = d.add_paragraph(); fr = foot.add_run("Source of truth: PROJECT-DNA.md + LOCKED_STATE.md. This brief is a derived rendering, regenerated when the DNA changes.")
    fr.italic=True; fr.font.size=Pt(8)
    p = OUT / "PlantMind_Executive_Brief_v1.docx"; d.save(p); return p

if __name__ == "__main__":
    x = excel(); b = docx()
    print("WROTE:", x.name, "+", b.name, "→", OUT)
